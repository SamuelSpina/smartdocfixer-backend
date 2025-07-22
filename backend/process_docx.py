# process_docx.py

import os
import tempfile
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches
from openai import OpenAI # <-- Import the new client

# Load environment variables from .env file for local development
load_dotenv()

# --- Initialize the OpenAI Client ---
# The client automatically looks for the OPENAI_API_KEY environment variable.
# This is the modern and recommended way.
try:
    client = OpenAI()
    # Test the connection by listing models (optional, but good for debugging)
    client.models.list() 
    print("Successfully connected to OpenAI.")
except Exception as e:
    print(f"Error: Could not initialize OpenAI client. Check your API key. Details: {e}")
    client = None

async def fix_document(file):
    """Fix grammar, clarity & formatting in a .docx upload using the modern OpenAI client."""
    
    # Check if the client was initialized successfully
    if not client:
        # If the client failed to init, we can't process the document.
        # This will raise an error that the user's request will see.
        raise ConnectionError("OpenAI service is not configured on the server.")

    # 1) Read and dump to temp file
    contents = await file.read()
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(contents)
        tmp_path = tmp.name

    # 2) Load doc
    doc = Document(tmp_path)
    print(f"Loaded {len(doc.paragraphs)} paragraphs…")

    # 3) AI system prompt - UPDATED
    system_msg = """You are SmartDocFixer AI, an expert editor. You improve grammar, clarity, and professional formatting. 
    Preserve essential structures like headings, lists, and tables. Your goal is to polish the text, not remove its core components.
    Do not add any conversational text or apologies like "Here is the fixed paragraph:".
    IMPORTANT: Your output must be plain text only. Do not use any Markdown formatting like **bold** or # headings. Just return the corrected text directly."""

    # 4) Iterate through paragraphs and fix them
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # Call OpenAI using the new client.chat.completions.create method
        try:
            # Note the new syntax here
            resp = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": system_msg},
                    {"role": "user",   "content": f"Correct and improve the following paragraph:\n\n{text}"}
                ],
                max_tokens=1024,
                temperature=0.2,
            )
            improved = resp.choices[0].message.content.strip()
        except Exception as e:
            # If the API call fails for any reason, log it and fall back to original text
            print(f"OpenAI API error on paragraph {i}:", e)
            improved = text  # Fallback to the original text

        # 5) Replace paragraph text
        para.text = improved

        # 6) Re-apply basic font formatting to ensure consistency
        for run in para.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)

    # 7) Set global document formatting
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 8) Save to a new file and clean up the temp file
    out_path = tmp_path.replace(".docx", "_fixed.docx")
    doc.save(out_path)
    os.unlink(tmp_path)
    print("Saved fixed document to:", out_path)
    return out_path
