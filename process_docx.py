from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openai
import os
import tempfile
from dotenv import load_dotenv

load_dotenv()

async def fix_document(file):
    """Main function to fix document grammar, clarity, and formatting"""
    
    # Read uploaded file
    contents = await file.read()
    
    # Create temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(contents)
        tmp_path = tmp.name
    
    # Load document
    doc = Document(tmp_path)
    
    print(f"Processing document with {len(doc.paragraphs)} paragraphs")
    
    # Fix each paragraph with AI
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():  # Only process non-empty paragraphs
            original_text = para.text
            print(f"Original paragraph {i}: {original_text[:100]}...")  # Show first 100 chars
            
            # Get AI-improved text
            improved_text = await improve_text_with_ai(original_text)
            print(f"Improved paragraph {i}: {improved_text[:100]}...")  # Show first 100 chars
            
            # Clear the paragraph and add improved text
            para.clear()
            run = para.add_run(improved_text)
            
            # Apply consistent formatting
            apply_formatting_to_run(run)
    
    # Apply document-wide formatting
    format_document(doc)
    
    # Save fixed document
    output_path = tmp_path.replace(".docx", "_fixed.docx")
    doc.save(output_path)
    
    print(f"Document processed successfully!")
    
    # Clean up original temp file
    os.unlink(tmp_path)
    
    return output_path

async def improve_text_with_ai(text):
    """Use OpenAI to improve grammar and clarity"""
    
    # Check if API key exists
    api_key = os.getenv("OPENAI_API_KEY")
    
    if not api_key:
        print("ERROR: No OPENAI_API_KEY found in environment variables")
        return f"[NO API KEY] {text}"
    
    if not api_key.startswith("sk-"):
        print(f"ERROR: Invalid API key format. Key starts with: {api_key[:10]}...")
        return f"[INVALID KEY] {text}"
    
    print(f"Using API key: {api_key[:10]}...{api_key[-10:]}")  # Show first/last 10 chars
    
    # Initialize client
    try:
        client = openai.OpenAI(api_key=api_key)
        print("OpenAI client initialized successfully")
    except Exception as e:
        print(f"ERROR creating OpenAI client: {e}")
        return f"[CLIENT ERROR] {text}"
    
    prompt = f"""Fix this text by improving grammar, clarity, and professionalism. Keep the same meaning but make it much better:

{text}

Return only the improved text."""
    
    try:
        print("Making API call to OpenAI...")
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a professional editor. Fix grammar, improve clarity, make text professional. Return only the improved text."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=800,
            temperature=0.2
        )
        
        print("API call successful!")
        improved = response.choices[0].message.content.strip()
        
        # Remove quotes if AI adds them
        if improved.startswith('"') and improved.endswith('"'):
            improved = improved[1:-1]
        if improved.startswith("'") and improved.endswith("'"):
            improved = improved[1:-1]
        
        return improved
    
    except Exception as e:
        print(f"DETAILED API ERROR: {type(e).__name__}: {str(e)}")
        return f"[API CALL FAILED: {str(e)}] {text}"

def apply_formatting_to_run(run):
    """Apply consistent formatting to a text run"""
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

def format_document(doc):
    """Apply document-wide formatting improvements"""
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Format potential headings
    for para in doc.paragraphs:
        text = para.text.strip()
        # Simple heuristic for headings: short lines without periods
        if text and len(text) < 60 and not text.endswith('.') and not text.endswith(','):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(14)
